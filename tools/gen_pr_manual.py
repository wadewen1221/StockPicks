#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成 PR 操作手册 DOCX
输出: E:\\技术文件资料\\PR操作手册_股票智能选股.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_chinese_font(run, size=11, bold=False, color=None):
    """设置中文字体（宋体/微软雅黑）+ 字号 + 颜色"""
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 中文字体需通过 rPr 元素设置
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_heading(doc, text, level=1):
    """添加标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    sizes = {0: 24, 1: 18, 2: 14, 3: 12}
    set_chinese_font(run, size=sizes.get(level, 12), bold=True, color=(0, 51, 102))
    p.style = doc.styles[f"Heading {min(level, 3)}"] if level <= 3 else doc.styles["Heading 3"]
    return p


def add_para(doc, text, size=11, bold=False, color=None, indent=0):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_chinese_font(run, size=size, bold=bold, color=color)
    return p


def add_step_box(doc, step_num, title, body):
    """添加步骤卡片（带边框感的段落）"""
    # 步骤标题
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(12)
    p1.paragraph_format.space_after = Pt(4)
    p1.paragraph_format.left_indent = Cm(0.3)
    run1 = p1.add_run(f"【步骤 {step_num}】 {title}")
    set_chinese_font(run1, size=12, bold=True, color=(204, 0, 0))

    # 步骤正文
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.6)
    p2.paragraph_format.space_after = Pt(8)
    for line in body:
        if line.startswith("📌"):
            run = p2.add_run("\n" + line)
            set_chinese_font(run, size=10, color=(0, 102, 51))
        elif line.startswith("⚠"):
            run = p2.add_run("\n" + line)
            set_chinese_font(run, size=10, color=(204, 102, 0))
        elif line.startswith("💻"):
            run = p2.add_run("\n" + line)
            set_chinese_font(run, size=10, color=(102, 102, 102), bold=False)
        else:
            run = p2.add_run(line)
            set_chinese_font(run, size=10)


def add_screenshot_placeholder(doc, description, expected_ui):
    """添加'截图占位符'——用文字描述界面布局 + 关键元素位置

    之所以不直接嵌图片:1) Vibe Coding 工作流中,文档主要由 AI 生成
    2) 截图需要实际操作才能获得 3) 文字描述足够清晰,用户操作时心里有数
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("📷 截图占位符 — ")
    set_chinese_font(run, size=10, bold=True, color=(153, 51, 102))
    run2 = p.add_run(description)
    set_chinese_font(run2, size=10, color=(102, 102, 102))

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.9)
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run("界面应该长这样:")
    set_chinese_font(run, size=10, bold=True)
    for line in expected_ui:
        p3 = doc.add_paragraph()
        p3.paragraph_format.left_indent = Cm(1.2)
        p3.paragraph_format.space_after = Pt(2)
        run = p3.add_run(line)
        set_chinese_font(run, size=9, color=(102, 102, 102))


def add_code_block(doc, code_text):
    """添加代码块(用等宽字体 + 灰底)"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    set_chinese_font(run, size=10, color=(0, 0, 153))
    run.font.name = "Consolas"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_table_simple(doc, headers, rows, col_widths=None):
    """添加简单表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    # 表头
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        set_chinese_font(run, size=10, bold=True, color=(255, 255, 255))
        # 单元格灰底
        tc_pr = hdr[i]._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "003366")
        tc_pr.append(shd)
    # 数据
    for r_idx, row in enumerate(rows, 1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            run = cells[c_idx].paragraphs[0].add_run(str(val))
            set_chinese_font(run, size=10)
    if col_widths:
        for col, w in zip(table.columns, col_widths):
            for cell in col.cells:
                cell.width = Cm(w)


# ============================================
# 主文档生成
# ============================================

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ================== 封面 ==================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(80)
run = title_p.add_run("GitHub PR 操作手册")
set_chinese_font(run, size=28, bold=True, color=(0, 51, 102))

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_p.paragraph_format.space_before = Pt(12)
run = subtitle_p.add_run("—— 股票智能选股 V2 项目")
set_chinese_font(run, size=14, color=(102, 102, 102))

# 项目元信息
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.paragraph_format.space_before = Pt(60)
for line in [
    "项目：stock-picks-v2",
    "适用读者：Vibe Coding 出身、首次接触 GitHub PR 的开发者",
    "阅读时间：约 25 分钟",
    "版本：v1.0（2026-07-21）",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    set_chinese_font(run, size=11, color=(102, 102, 102))

doc.add_page_break()

# ================== 目录 ==================
add_heading(doc, "目录", 1)
toc_items = [
    ("1. 什么是 PR（Pull Request）?", "1"),
    ("2. 准备工作：装 Git + 注册 GitHub", "2"),
    ("3. 第一次：Fork 别人的项目并提 PR", "3"),
    ("4. 在自己项目里提 PR（团队协作）", "5"),
    ("5. PR 评审怎么进行", "7"),
    ("6. PR 的生命周期（看 GitHub Actions）", "8"),
    ("7. 常见问题排查", "9"),
    ("附录 A. Git 常用命令速查", "10"),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(item)
    set_chinese_font(run, size=11)
    run2 = p.add_run("\t" + "P." + page)
    set_chinese_font(run2, size=10, color=(102, 102, 102))

doc.add_page_break()

# ================== 第 1 章 ==================
add_heading(doc, "1. 什么是 PR（Pull Request）?", 1)
add_para(
    doc,
    "PR = Pull Request = 「合并申请」。"
    "你修改了别人（或自己）的代码,想把这些修改「合入」主代码库,就需要提一个 PR。",
    size=11,
)
add_para(doc, "一句话类比：", bold=True, size=12)
add_para(
    doc,
    "上市公司要搞重大事项(增发/并购),需要向董事会提交提案,经独立董事审议,才能执行。\n"
    "代码要合入主分支,需要提一个 PR,经其他开发者 + 自动测试审查,才能合入。",
    size=11,
    indent=0.3,
)

add_heading(doc, "为什么需要 PR?", 2)
add_para(doc, '防止"一言堂"——任意开发者都能改主代码 = 灾难。', size=11)
add_para(doc, "多人协作——你的修改要让团队成员看到、讨论、把关。", size=11)
add_para(doc, "留下记录——所有讨论、修改、决策都有案可查。", size=11)
add_para(doc, "触发自动化——CI/CD 会自动跑测试,确保你的改动没把代码改坏。", size=11)

doc.add_page_break()

# ================== 第 2 章 ==================
add_heading(doc, "2. 准备工作：装 Git + 注册 GitHub", 1)
add_para(doc, "如果你已经有 GitHub 账号 + 本地装了 Git,跳过本节。", size=11, color=(102, 102, 102))

add_step_box(
    doc, 1, "安装 Git",
    [
        "Windows:访问 https://git-scm.com/download/win,下载 64-bit 安装包",
        "一路 Next 即可(默认选项已合理)",
        "💻 安装完打开 PowerShell,输入 git --version 应看到 git version 2.x.x",
    ],
)
add_step_box(
    doc, 2, "注册 GitHub 账号",
    [
        "访问 https://github.com/signup",
        "填邮箱、密码、用户名(用户名会成为您的公开 ID,如 longdashen)",
        "📌 重要:绑定的邮箱要能收到验证邮件(用来收 PR 评论通知)",
    ],
)
add_step_box(
    doc, 3, "配置 Git 身份",
    [
        "打开 PowerShell,执行(把名字和邮箱换成你自己的):",
        "💻 git config --global user.name \"你的名字\"",
        "💻 git config --global user.email \"你的邮箱@xxx.com\"",
        "📌 这两个信息会出现在你每次 commit 的记录里,务必写真实信息",
    ],
)
add_step_box(
    doc, 4, "把 stock-picks-v2 项目 clone 到本地",
    [
        "本项目已经在您本地 D:\\\\stock-picks-v2\\\\,无需重新 clone",
        "如果是别人的开源项目,命令是:",
        "💻 git clone https://github.com/用户名/stock-picks-v2.git",
        "💻 cd stock-picks-v2",
    ],
)

doc.add_page_break()

# ================== 第 3 章 ==================
add_heading(doc, "3. 第一次:Fork 别人的项目并提 PR", 1)
add_para(
    doc,
    "如果你要给别人的开源项目(stock-picks-v2)提代码贡献,流程是 Fork → 改 → 提 PR。",
    size=11,
)

add_step_box(
    doc, 1, "在 GitHub 上 Fork 项目",
    [
        "打开项目主页:https://github.com/原作者/stock-picks-v2",
        "右上角点 [Fork] 按钮(大蓝色按钮)",
        "几秒后,GitHub 会在你账号下创建一个副本:https://github.com/你的用户名/stock-picks-v2",
    ],
)
add_screenshot_placeholder(
    doc,
    "Fork 按钮位置",
    [
        "[ 顶部右侧区域 ]",
        "Watch  [Star]  [Fork]   ← 点这个 Fork",
        "",
        "[ Fork 完成后的跳转页面 ]",
        "中间显示:yourname/stock-picks-v2 was created",
        "几秒后自动跳转到 https://github.com/你的用户名/stock-picks-v2",
    ],
)

add_step_box(
    doc, 2, "把 Fork 后的项目 clone 到本地",
    [
        "在 PowerShell 里:",
        "💻 git clone https://github.com/你的用户名/stock-picks-v2.git",
        "💻 cd stock-picks-v2",
        "💻 git remote add upstream https://github.com/原作者/stock-picks-v2.git",
        "📌 upstream 是「上游」,用来同步原作者的更新",
    ],
)

add_step_box(
    doc, 3, "新建一个分支改代码",
    [
        "千万不要直接在 master 上改!一定要新建分支:",
        "💻 git checkout -b fix-rsrs-bug",
        "💻 ... 修改代码 ...",
        "💻 git add .",
        "💻 git commit -m \"修复 RSRS 斜率计算错误\"",
        "💻 git push origin fix-rsrs-bug",
    ],
)

add_step_box(
    doc, 4, "在 GitHub 上提 PR",
    [
        "回到你的 Fork 页面 https://github.com/你的用户名/stock-picks-v2",
        "你会看到黄色横幅:\"fix-rsrs-bug had recent pushes · Open PR\"",
        "点 [Compare & pull request] 按钮",
        "",
        "⚠ 如果没看到黄色横幅:",
        "  → 手动点 [Pull requests] 标签页 → [New pull request]",
        "  → 左侧选 yourname:fix-rsrs-bug",
        "  → 右侧选 原作者:master",
    ],
)
add_screenshot_placeholder(
    doc,
    "PR 创建页面布局",
    [
        "[ 顶部 ]",
        "Open pull request",
        "",
        "[ 左侧 base repository]",
        "base: 原作者/stock-picks-v2  <-  base: master",
        "",
        "[ 右侧 head repository]",
        "head: 你的用户名/stock-picks-v2 <- compare: fix-rsrs-bug",
        "",
        "[ Able to merge 绿色提示 = 没有冲突,可以合并 ]",
    ],
)

add_step_box(
    doc, 5, "填写 PR 标题和说明",
    [
        "Title: 简短说明这次改了什么,如 \"修复 RSRS 斜率永远为 0 的 bug\"",
        "Description: 详细说明,推荐模板:",
        "💻 ## 改了什么",
        "💻 - 把 calculate_rsrs() 的空数据分支加上 slope 字段",
        "💻 ",
        "💻 ## 怎么测的",
        "💻 - pytest 87 个测试全过",
        "💻 - 真实股票 000001 验证 slope=0.9458",
        "",
        "📌 写清楚\"为什么改\"和\"怎么验证\"会让 reviewer 更容易通过",
    ],
)
add_step_box(
    doc, 6, "提交 PR,等待 review",
    [
        "点 [Create pull request] 按钮",
        "页面会跳转到 PR 详情页",
        "等原作者/维护者评审 — 通常 1-3 天",
        "📌 如果 reviewer 提了评论(行内代码评论),你直接在网页上回复+push 新 commit,PR 会自动更新",
    ],
)

doc.add_page_break()

# ================== 第 4 章 ==================
add_heading(doc, "4. 在自己项目里提 PR(团队协作)", 1)
add_para(
    doc,
    "如果你在自己团队的 stock-picks-v2 项目上协作,流程类似,但不需要 Fork — "
    "直接在主项目里建分支、改、提 PR 即可。",
    size=11,
)

add_step_box(
    doc, 1, "先 clone 团队主项目",
    [
        "💻 git clone https://github.com/你的团队/stock-picks-v2.git",
        "💻 cd stock-picks-v2",
    ],
)

add_step_box(
    doc, 2, "保持本地 master 与远程同步",
    [
        "💻 git checkout master",
        "💻 git pull origin master",
        "📌 每次开工前先 pull,避免和别人冲突",
    ],
)

add_step_box(
    doc, 3, "建分支 + 改代码 + 提交",
    [
        "💻 git checkout -b feature/new-indicator",
        "💻 ... 修改代码,跑测试:pytest tests/ ...",
        "💻 git add .",
        "💻 git commit -m \"feat: 添加布林带指标打分\"",
        "💻 git push origin feature/new-indicator",
    ],
)
add_screenshot_placeholder(
    doc,
    "分支命名建议",
    [
        "[ 好的命名 ]",
        "  fix-rsrs-slope-bug          ← 修 bug",
        "  feature/add-docker-compose   ← 加新功能",
        "  refactor/split-selector      ← 重构",
        "  docs/update-pr-manual        ← 文档",
        "  test/add-rsrs-cases          ← 测试",
        "",
        "[ 不好的命名 ]",
        "  test       ← 看不出要干嘛",
        "  my_branch  ← 名字太模糊",
        "  fix        ← fix 啥?",
    ],
)

add_step_box(
    doc, 4, "在 GitHub 上提 PR",
    [
        "1. 打开 https://github.com/你的团队/stock-picks-v2",
        "2. 你会看到黄色横幅:\"feature/new-indicator had recent pushes\"",
        "3. 点 [Compare & pull request]",
        "4. base 选 master,compare 选 feature/new-indicator",
        "5. 写 Title + Description,点 [Create pull request]",
    ],
)

doc.add_page_break()

# ================== 第 5 章 ==================
add_heading(doc, "5. PR 评审怎么进行", 1)
add_para(doc, "PR 提了之后,会发生这些事:", size=11, bold=True)

add_table_simple(
    doc,
    ["角色", "做什么", "你能看到什么"],
    [
        ["🤖 GitHub Actions", "自动跑 pytest + npm build + flake8 等", "绿色 ✓ 或红色 ✗,看 .github/workflows/"],
        ["👀 团队成员", "读 diff,在行内评论,提建议", "PR 页面 Conversation 标签 + Files Changed"],
        ["👑 维护者", "Approve / Request Changes / Merge", "Merge 按钮亮起 = 可合并"],
    ],
    col_widths=[3, 6, 6],
)

add_para(doc, "评论的常见位置:", size=11, bold=True, color=(0, 51, 102))
add_screenshot_placeholder(
    doc,
    "PR 评论界面布局",
    [
        "[ Files Changed 标签页 ]",
        "  src/main.py",
        "    @@ -10,7 +10,7 @@",
        "    -    old_code   ← 红色,被删除",
        "    +    new_code   ← 绿色,新增",
        "         (鼠标悬停 → 出现 + 号 → 可加行内评论)",
        "",
        "[ 评论示例 ]",
        "  💬 龙大神 commented 2 hours ago",
        "  这里为什么用 pandas 而不是 numpy?",
        "  [ Reply ] [ React ]",
    ],
)

add_step_box(
    doc, 1, "如何响应评论",
    [
        "在网页上点 [Reply] 直接回复",
        "如果你改了代码:",
        "  → 在本地改 → git add . → git commit -m \"address review comment\" → git push",
        "  → PR 会自动更新,新 commit 会出现在 PR 页面",
        "",
        "📌 保持礼貌、就事论事、引用对方的评论,不要发火",
    ],
)

doc.add_page_break()

# ================== 第 6 章 ==================
add_heading(doc, "6. PR 的生命周期(看 GitHub Actions)", 1)
add_para(
    doc,
    "每次 push 新 commit,GitHub Actions 会自动跑测试套件。这就是 CI 的实际表现。",
    size=11,
)
add_screenshot_placeholder(
    doc,
    "PR 页面底部 Checks 区域",
    [
        "[ PR 页面 → 底部 ]",
        "  ⌄ Checks",
        "    ✓ Backend (Python 3.11, ubuntu-latest)  - 1m 23s",
        "    ✓ Backend (Python 3.12, ubuntu-latest)  - 1m 18s",
        "    ✓ Frontend (Node 20, ubuntu-latest)      - 0m 48s",
        "    ✓ Lint (Python)                          - 0m 12s",
        "",
        "[ 全部绿色 = 可以合并 ]",
        "[ 任何一个红色 ✗ = 禁止合并,先修测试 ]",
    ],
)

add_step_box(
    doc, 1, "点击查看详细日志",
    [
        "点任意 Check 名称(如 Backend ... Python 3.12 ...)",
        "进入 logs 页面,可以看到:",
        "  - pytest 跑了哪些测试(每个 ✓/✗)",
        "  - 失败时的 traceback",
        "  - 覆盖率报告",
    ],
)

add_step_box(
    doc, 2, "如果 CI 红了",
    [
        "📌 常见原因:",
        "  1. 你的代码改坏了某个测试 → 本地跑 pytest 找出来,修代码",
        "  2. 缺依赖 → 检查 requirements.txt 是否加全",
        "  3. 跨平台问题 → 在 Windows 写的换行符跑到 Linux 失败(用 git config --global core.autocrlf input)",
        "  4. 偶发网络失败 → 重新 push 一次",
    ],
)

doc.add_page_break()

# ================== 第 7 章 ==================
add_heading(doc, "7. 常见问题排查", 1)

# 表格: 错误 → 排查
add_para(doc, "Q1: git push 时说\"Permission denied\"", bold=True, size=12, color=(204, 0, 0))
add_para(
    doc,
    "A: 你的 GitHub 账号没认证。两种解决:",
    size=11,
    indent=0.3,
)
add_code_block(
    doc,
    "方案 A(简单):用 HTTPS 协议 + Personal Access Token\n"
    "  1. 打开 GitHub → Settings → Developer settings → Personal access tokens\n"
    "  2. Generate new token (classic),勾选 repo 权限\n"
    "  3. 复制 token,git push 时密码填这个 token\n"
    "\n"
    "方案 B(推荐):用 SSH 协议\n"
    "  1. ssh-keygen -t rsa -b 4096 -C \"你的邮箱\"\n"
    "  2. 把 ~/.ssh/id_rsa.pub 内容贴到 GitHub → Settings → SSH and GPG keys\n"
    "  3. git remote set-url origin git@github.com:用户名/stock-picks-v2.git",
)

add_para(doc, "Q2: PR 提了但说\"Can't automatically merge\"", bold=True, size=12, color=(204, 0, 0))
add_para(
    doc,
    "A: 你的分支落后于 master,有冲突。两种解决:",
    size=11,
    indent=0.3,
)
add_code_block(
    doc,
    "方案 A:网页解决(简单)\n"
    "  → PR 页面会显示\"Resolve conflicts\"按钮,点开在网页上点区块选择保留哪边\n"
    "\n"
    "方案 B:本地解决(灵活)\n"
    "  git checkout master && git pull origin master\n"
    "  git checkout 你的分支 && git merge master\n"
    "  # 打开冲突文件,删掉 <<<<<< HEAD 之类的标记\n"
    "  git add . && git commit -m \"merge master\" && git push",
)

add_para(doc, "Q3: 我 commit 后想改 message", bold=True, size=12, color=(204, 0, 0))
add_para(
    doc,
    "A:git commit --amend -m \"新 message\"。⚠ 只对最新 commit 有用,且没 push 之前。",
    size=11,
    indent=0.3,
)

add_para(doc, "Q4: 我提交错了东西,想撤回", bold=True, size=12, color=(204, 0, 0))
add_para(
    doc,
    "A:看你是要撤回哪个:",
    size=11,
    indent=0.3,
)
add_code_block(
    doc,
    "撤回最近 1 个 commit(保留修改):  git reset --soft HEAD~1\n"
    "撤回最近 1 个 commit(丢弃修改):    git reset --hard HEAD~1   ⚠ 慎用\n"
    "撤回已 push 的 commit(在 PR 里):   git revert HEAD,再 push,会新增一个反向 commit",
)

add_para(doc, "Q5: 我想看某个文件的历史", bold=True, size=12, color=(204, 0, 0))
add_para(
    doc,
    "A:git log --follow -- 文件名。",
    size=11,
    indent=0.3,
)

doc.add_page_break()

# ================== 附录 A ==================
add_heading(doc, "附录 A. Git 常用命令速查", 1)

add_para(doc, "一、新人第一天", bold=True, size=12, color=(0, 51, 102))
add_code_block(
    doc,
    "git clone <url>                # 克隆项目到本地\n"
    "cd stock-picks-v2\n"
    "git status                     # 看当前状态(改了啥)\n"
    "git log --oneline              # 看最近 10 次提交(一行一个)\n"
    "git branch                     # 看本地所有分支",
)

add_para(doc, "二、修改 + 提交", bold=True, size=12, color=(0, 51, 102))
add_code_block(
    doc,
    "git checkout -b feature/xxx    # 新建并切换分支\n"
    "... 改代码 ...\n"
    "git diff                       # 看修改了啥(还没 add)\n"
    "git add .                      # 把所有修改加入暂存区\n"
    "git add -p                     # 交互式 add(可逐块选择)\n"
    "git commit -m \"message\"        # 提交\n"
    "git push origin feature/xxx    # 推送到远程",
)

add_para(doc, "三、同步远程", bold=True, size=12, color=(0, 51, 102))
add_code_block(
    doc,
    "git fetch origin               # 拉取远程所有更新(不合并)\n"
    "git pull origin master         # 拉取并合并 master 到当前分支\n"
    "git remote -v                  # 看远程仓库地址",
)

add_para(doc, "四、紧急救助", bold=True, size=12, color=(0, 51, 102))
add_code_block(
    doc,
    "git stash                      # 暂存当前修改(可切走)\n"
    "git stash pop                  # 恢复暂存\n"
    "git reset --hard HEAD~1        # 撤销最近 1 个 commit(丢弃修改) ⚠\n"
    "git revert <commit-hash>       # 反转某次 commit 的修改\n"
    "git reflog                     # 看所有 HEAD 移动记录(救命用)",
)

# ================== 收尾 ==================
doc.add_page_break()
add_heading(doc, "文档结束", 1)
add_para(
    doc,
    "本手册基于 stock-picks-v2 项目实战编写,适用于 Vibe Coding 出身、"
    "首次接触 GitHub 协作流程的开发者。",
    size=11,
)
add_para(
    doc,
    "如有疑问,提 Issue 或在 PR 评论区讨论。",
    size=11,
)
add_para(
    doc,
    "—— 龙大神 编制 / 2026-07-21",
    size=10,
    color=(102, 102, 102),
)

# ================== 落盘 ==================
output_dir = Path(r"E:\技术文件资料")
output_dir.mkdir(parents=True, exist_ok=True)
output_path = output_dir / "PR操作手册_股票智能选股.docx"
doc.save(str(output_path))

print(f"✅ 已生成:{output_path}")
print(f"   文件大小:{output_path.stat().st_size / 1024:.1f} KB")
print(f"   页数估算:{output_path.stat().st_size // 5000} 页(粗略)")